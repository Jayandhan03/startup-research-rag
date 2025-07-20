from docx import Document
from datetime import datetime

def render_to_docx(query: str, response: str, filename: str = "output.docx") -> None:
    """
    Renders the query and response into a formatted .docx file.

    Args:
        query (str): The original user query.
        response (str): The LLM-generated answer.
        filename (str): Output filename.
    """
    doc = Document()
    doc.add_heading("Startup Investment Research Report", level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Query: {query}")
    doc.add_heading("AI-Generated Answer", level=2)
    doc.add_paragraph(response)
    doc.save(filename)
